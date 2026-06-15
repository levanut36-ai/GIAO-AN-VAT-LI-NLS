import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# Cấu hình giao diện ứng dụng rộng rãi, chuyên nghiệp
st.set_page_config(page_title="Giáo Án Vật Lí 5512 & NLS", page_icon="⚛️", layout="wide")

st.markdown("<h2 style='text-align: center; color: #1E88E5;'>⚛️ ỨNG DỤNG SOẠN GIÁO ÁN VẬT LÍ TRỌN GÓI MỘT MẠCH (CV 5512 & TT 02/2025)</h2>", unsafe_allow_html=True)
st.write("---")

# Thanh cấu hình bên trái
st.sidebar.header("🔑 Cấu hình hệ thống")
api_key = st.sidebar.text_input("Nhập Gemini API Key của bạn:", type="password", help="Lấy mã này từ Google AI Studio")

st.sidebar.write("---")
st.sidebar.header("📚 Năng lực số (TT 02/2025)")
mien_nls = st.sidebar.selectbox(
    "Chọn miền năng lực số lồng ghép chủ đạo:",
    [
        "Miền I: Khai thác dữ liệu và thông tin trong môi trường số",
        "Miền II: Giao tiếp và hợp tác trong môi trường số",
        "Miền III: Sáng tạo nội dung số",
        "Miền IV: An toàn và bảo mật thông tin",
        "Miền V: Giải quyết sự cố công nghệ và vấn đề kỹ thuật",
        "Miền VI: Ứng dụng Trí tuệ nhân tạo (AI) vào học tập"
    ]
)

# Hàm thiết lập khoảng cách ô trong file Word
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Bộ xử lý bóc tách văn bản để vẽ bảng 3 cột tự động cho TỪNG hoạt động trong file Word
def tao_file_word_tron_goi(text_content):
    doc = Document()
    
    # Định dạng phông chữ chuẩn văn bản hành chính Việt Nam
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Chia tách văn bản thành các khối Hoạt động lớn độc lập
    if '**Hoạt động' in text_content:
        parts = text_content.split('**Hoạt động')
        header_part = parts[0]
        activity_parts = parts[1:]
    else:
        header_part = text_content
        activity_parts = []
    
    # Khối 1: Ghi phần Khung đầu, Mục I, Mục II trước khi vào tiến trình hoạt động
    header_lines = header_part.split('\n')
    for line in header_lines:
        line_clean = line.replace('**', '').replace('###', '').replace('##', '').replace('#', '').strip()
        if not line_clean:
            continue
        p = doc.add_paragraph()
        if line.startswith('#') or any(m in line_clean for m in ["I. MỤC TIÊU", "II. THIẾT BỊ", "III. TIẾN TRÌNH", "1.", "2.", "3."]):
            p.add_run(line_clean).bold = True
        else:
            p.add_run(line_clean)
            
    # Xử lý bóc tách và kẻ bảng 3 cột tự động từ Hoạt động 1 đến Hoạt động 4
    for part in activity_parts:
        full_activity_text = "**Hoạt động" + part
        lines = full_activity_text.split('\n')
        
        buoc_hd = {"Bước 1": "", "Bước 2": "", "Bước 3": "", "Bước 4": ""}
        buoc_nls = {"Bước 1": "Không tích hợp.", "Bước 2": "Không tích hợp.", "Bước 3": "Không tích hợp.", "Bước 4": "Không tích hợp."}
        current_buoc = None
        has_d_section = False
        
        for line in lines:
            line_clean = line.replace('**', '').replace('###', '').replace('##', '').replace('#', '').strip()
            if not line_clean:
                continue
                
            if any(b in line_clean for b in ["Bước 1:", "Chuyển giao nhiệm vụ"]):
                current_buoc = "Bước 1"
                has_d_section = True
            elif any(b in line_clean for b in ["Bước 2:", "Thực hiện nhiệm vụ"]):
                current_buoc = "Bước 2"
            elif any(b in line_clean for b in ["Bước 3:", "Báo cáo, thảo luận"]):
                current_buoc = "Bước 3"
            elif any(b in line_clean for b in ["Bước 4:", "Kết luận, nhận định"]):
                current_buoc = "Bước 4"
            elif "IV. KIỂM TRA" in line_clean:
                current_buoc = None
                has_d_section = False
                p = doc.add_paragraph()
                p.add_run(line_clean).bold = True
                continue
                
            if not has_d_section:
                p = doc.add_paragraph()
                if line.startswith('#') or any(m in line_clean for m in ["Hoạt động", "a)", "b)", "c)", "d)"]):
                    p.add_run(line_clean).bold = True
                else:
                    p.add_run(line_clean)
            else:
                if current_buoc:
                    if "Tích hợp NLS:" in line_clean or "Năng lực số tích hợp:" in line_clean:
                        clean_nls = line_clean.replace("Tích hợp NLS:", "").replace("Năng lực số tích hợp:", "").strip()
                        buoc_nls[current_buoc] = clean_nls if clean_nls else "Không tích hợp."
                    else:
                        buoc_hd[current_buoc] += line_clean + "\n"
                        
        if has_d_section:
            p_title = doc.add_paragraph()
            p_title.add_run("d) Tổ chức thực hiện:").bold = True
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows.cells
            hdr_cells[0].text = 'Tiến trình sư phạm'
            hdr_cells[1].text = 'Hoạt động cụ thể của Giáo viên và Học sinh'
            hdr_cells[2].text = 'Năng lực số'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                
            mapping = [
                ("Bước 1: Chuyển giao nhiệm vụ", "Bước 1"),
                ("Bước 2: Thực hiện nhiệm vụ", "Bước 2"),
                ("Bước 3: Báo cáo, thảo luận", "Bước 3"),
                ("Bước 4: Kết luận, nhận định", "Bước 4")
            ]
            
            for ten_buoc, key in mapping:
                row_cells = table.add_row().cells
                row_cells[0].text = ten_buoc
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                
                noi_dung_buoc = buoc_hd[key].replace(ten_buoc, "").strip()
                row_cells[1].text = noi_dung_buoc if noi_dung_buoc else "Giáo viên điều phối hoạt động giảng dạy theo phân phối chương trình."
                row_cells[2].text = buoc_nls[key]
                
                row_cells[0].width = Inches(1.5)
                row_cells[1].width = Inches(4.0)
                row_cells[2].width = Inches(1.5)
                set_cell_margins(row_cells[0])
                set_cell_margins(row_cells[1])
                set_cell_margins(row_cells[2])
            doc.add_paragraph() # Khoảng cách trống sau bảng
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# Khung nhập thông tin bài học Vật lí ở giao diện chính
st.subheader("📝 Thông tin chung bài học")
col1, col2 = st.columns(2)

with col1:
    truong = st.text_input("Tên trường / Cơ sở giáo dục:", placeholder="Ví dụ: THPT Nguyễn Du...")
    giao_vien = st.text_input("Họ và tên giáo viên:", placeholder="Ví dụ: Nguyễn Văn A...")
    khoi_lop = st.selectbox("Khối lớp (Môn Vật lí - Sách Kết nối tri thức):", ["Vật lí 10", "Vật lí 11", "Vật lí 12"])

with col2:
    to_chuyen_mon = st.text_input("Tổ / Nhóm chuyên môn:", placeholder="Ví dụ: Tổ Vật lí - Công nghệ...")
    ten_bai = st.text_input("Tên bài dạy chính xác:", placeholder="Ví dụ: Động lượng và định luật bảo toàn động lượng...")
    so_tiet = st.number_input("Thời gian thực hiện (Số tiết):", min_value=1, max_value=6, value=2)

muc_tieu_rieng = st.text_area("Yêu cầu bổ sung hoặc công cụ số muốn sử dụng (Nếu có):", 
                              placeholder="Ví dụ: Sử dụng thí nghiệm ảo PhET, điện thoại thông minh...")

if st.button("🚀 BẮT ĐẦU SOẠN GIÁO ÁN TRỌN GÓI MỘT MẠCH", type="primary", use_container_width=True):
    if not api_key:
        st.error("❌ Vui lòng nhập API Key ở thanh menu bên trái trước!")
    elif not ten_bai:
        st.warning("⚠️ Vui lòng nhập Tên bài dạy chính xác.")
    else:
        with st.spinner("⏳ Trợ lý AI đang áp dụng siêu dung lượng để soạn một mạch toàn bộ giáo án chuẩn... Vui lòng đợi trong giây lát!"):
            try:
                genai.configure(api_key=api_key)
                # Kích hoạt toàn bộ sức mạnh số từ đầu ra cho Gemini 2.5 Flash
                model_config = genai.GenerationConfig(max_output_tokens=65535, temperature=0.3)
                model = genai.GenerativeModel('gemini-2.5-flash', generation_config=model_config)
                
                tong_thoi_gian = so_tiet * 45
                
                # Sử dụng cấu trúc Triple Quotes thuần túy không lồng ngoặc để bảo đảm an toàn cú pháp 100%
                prompt = f"""
                Bạn là một chuyên gia giáo dục Vật lí xuất sắc cốt cán. Hãy thiết kế một kế hoạch bài dạy (Giáo án) HOÀN CHỈNH, TRỌN GÓI TỪ ĐẦU ĐẾN CUỐI bám sát chương trình SGK KẾT NỐI TRI THỨC môn Vật lí.
                - Bài dạy: {ten_bai}
                - Khối: {khoi_lop}
                - Số tiết: {so_tiet} tiết (Tổng cộng {tong_thoi_gian} phút).
                - Trường: {truong}
                - Giáo viên: {giao_vien}
                - Tổ: {to_chuyen_mon}
                - Định hướng lồng ghép Năng lực số: {mien_nls} bám sát Thông tư 02/2025/TT-BGDĐT.
                - Yêu cầu riêng: {muc_tieu_rieng}

                YÊU CẦU CẤU TRÚC: Hãy viết một mạch liên tục đầy đủ tất cả các mục lớn từ đầu đến cuối bài (từ mục I cho đến hết mục IV), không được dừng lại giữa chừng, không dùng dấu ba chấm để bỏ lửng nội dung. Cấu trúc gồm:
                **I. MỤC TIÊU** (Kiến thức; Năng lực chung, đặc thù, Năng lực số thành phần cụ thể; Phẩm chất).
